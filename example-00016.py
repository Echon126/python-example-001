from docx import Document

doc = Document()
doc.add_paragraph(u'对中国能源问题的思考', 'Title')
doc.add_paragraph(u'作者', 'Subtitle')
doc.add_paragraph(u'摘要：本文阐明了能源问题的重要性...', 'Body Text 2')
doc.add_paragraph(u'能源问题的重要性', 'Heading 1')
doc.add_paragraph(u'人类的能源利用经历了从薪柴时代到煤炭时代...')
doc.add_paragraph(u'能源是现代经济社会发展的基础', 'Heading 2')
doc.add_paragraph(u'现代经济社会发展建立在高水平物质文明')
p = doc.add_paragraph(u'能源是经济社会发展的重要制约因素')
p.style = 'Heading 1'
doc.save('demo.docx')
