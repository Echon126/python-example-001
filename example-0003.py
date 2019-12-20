## docx word 文档解析读取
from docx import Document


# 获取文档中的表格
# doc.tables
# 读取第一个表格
# tb1 = doc.tables[0]
# 获取第一个表格的行
# tb1.rows
# 读取表格第一行的单元格
# row_cells = tb1.row[0].cells

# 读取第一行所有单元格的内容
# for cell in row_cells:
#     print(cell.text)
def with_docx_word(path):
    # 读取文档
    doc = Document(path)
    tables = doc.tables
    for i in range(len(tables)):
        tb = tables[i]
        tb_rows = tb.rows
        for i in range(len(tb_rows)):
            row_data = []
            row_cells = tb_rows[i].cells
            for cell in row_cells:
                row_data.append(cell.text)

            print(row_data)


if __name__ == '__main__':
    with_docx_word(r"E:\test\test.docx")
